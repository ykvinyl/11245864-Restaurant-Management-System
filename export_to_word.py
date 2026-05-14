import os
import platform
import sys
import io
from docx import Document
from docx.shared import Pt

# Ép hệ thống sử dụng UTF-8 để in ra màn hình không bị lỗi Unicode
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def export_project_to_docx(project_dir, output_file="Project_Export.docx"):
    doc = Document()
    doc.add_heading('Project Skeleton & Source Code', 0)

    # Loại bỏ các thư mục rác để file Word không bị quá nặng
    exclude_dirs = {'venv', '.git', '__pycache__', 'migrations', 'node_modules', '.idea', '.vscode'}
    allowed_extensions = {'.py', '.html', '.css', '.js', '.txt', '.md', '.sql'}

    # 1. TẠO SKELETON (CÂY THƯ MỤC)
    doc.add_heading('1. Directory Structure', level=1)
    tree_structure = []
    
    root_basename = os.path.basename(os.path.abspath(project_dir)) or "Project_Root"
    tree_structure.append(f"{root_basename}/")
    
    for root, dirs, files in os.walk(project_dir):
        dirs[:] = [d for d in dirs if d not in exclude_dirs]
        
        rel_path = os.path.relpath(root, project_dir)
        level = 0 if rel_path == '.' else rel_path.count(os.sep) + 1
            
        if rel_path != '.':
            indent = '    ' * level
            tree_structure.append(f'{indent}{os.path.basename(root)}/')
        
        subindent = '    ' * (level + 1) if rel_path != '.' else '    '
        for f in files:
            if any(f.endswith(ext) for ext in allowed_extensions):
                tree_structure.append(f'{subindent}{f}')

    doc.add_paragraph('\n'.join(tree_structure))

    # 2. XUẤT NỘI DUNG CODE
    doc.add_heading('2. Source Code', level=1)
    
    for root, dirs, files in os.walk(project_dir):
        dirs[:] = [d for d in dirs if d not in exclude_dirs]
        for file in files:
            if any(file.endswith(ext) for ext in allowed_extensions):
                file_path = os.path.join(root, file)
                rel_file_path = os.path.relpath(file_path, project_dir)
                
                doc.add_heading(f"File: {rel_file_path}", level=2)
                
                try:
                    # Đọc file với utf-8, bỏ qua ký tự lỗi nếu có
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                        content = f.read()
                        p = doc.add_paragraph(content)
                        for run in p.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                except Exception as e:
                    doc.add_paragraph(f"[Lỗi đọc file: {e}]")

    try:
        doc.save(output_file)
        print(f"SUCCESS: Da xuat file {output_file}")
    except PermissionError:
        print(f"ERROR: Vui lòng đóng file {output_file} đang mở trong Word!")
    except Exception as e:
        print(f"ERROR: Khong the luu file vi: {e}")

if __name__ == "__main__":
    # Đảm bảo bạn đang ở đúng thư mục dự án
    target_dir = '.' 
    export_project_to_docx(target_dir)
    
    output_path = "Project_Export.docx"
    if os.path.exists(output_path):
        if platform.system() == "Windows":
            os.startfile(output_path)
        elif platform.system() == "Darwin":
            os.system(f'open "{output_path}"')